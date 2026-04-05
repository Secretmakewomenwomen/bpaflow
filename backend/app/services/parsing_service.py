from __future__ import annotations

# BytesIO 用于把二进制文件内容包装成类文件对象，交给第三方解析库读取。
from io import BytesIO

# Settings 提供 OCR 语言等解析阶段需要的配置。
from app.core.config import Settings
# ParsedDocument 和 ParsedSegment 是解析层统一输出的标准数据结构。
from app.services.document_types import ParsedDocument, ParsedSegment


class ParsingService:
    # 解析服务只负责“把文件转成标准文本片段”，不在这一层做清洗和切块，职责保持单一。
    def __init__(self, settings: Settings) -> None:
        # 保存配置，后面 OCR 语言等参数都从这里读取。
        self.settings = settings

    # 按文件类型路由到不同解析器，统一产出 ParsedDocument，方便后续流水线复用。
    def parse(
        self,
        filename: str,
        file_ext: str,
        mime_type: str,
        content: bytes,
    ) -> ParsedDocument:
        # 统一把扩展名转成小写并去掉前导点，便于后续分支判断。
        extension = file_ext.lower().lstrip(".")
        # docx 文件走 Word 正文/表格解析逻辑。
        if extension == "docx":
            return self._parse_docx(content)
        # pdf 文件优先抽文本，缺失页再 OCR 兜底。
        if extension == "pdf":
            return self._parse_pdf(content)
        # png 文件直接走 OCR，提取图片里的文字。
        if extension == "png":
            return self._parse_png(content)
        # 其他格式当前不支持向量化，直接抛出明确异常。
        raise ValueError(f"Unsupported vectorization file type: {filename}")

    # docx 先按正文段落抽取，保留自然语义边界，便于后续切块时少打断句子。
    def _parse_docx(self, content: bytes) -> ParsedDocument:
        try:
            # Document 是 python-docx 的主入口，用于打开 docx。
            from docx import Document
            # DocxDocument 用于判断当前对象是不是文档根节点。
            from docx.document import Document as DocxDocument
            # CT_Tbl 表示 docx 底层 XML 里的表格节点类型。
            from docx.oxml.table import CT_Tbl
            # CT_P 表示 docx 底层 XML 里的段落节点类型。
            from docx.oxml.text.paragraph import CT_P
            # Table 是高层表格包装对象。
            from docx.table import Table
            # Paragraph 是高层段落包装对象。
            from docx.text.paragraph import Paragraph
        except ImportError as exc:
            # 如果 python-docx 没装，给出更聚焦的运行时异常。
            raise RuntimeError("python-docx is not installed") from exc

        # 把字节内容包装成文件流并打开成 docx 文档对象。
        document = Document(BytesIO(content))
        # 收集解析出的标准片段。
        segments: list[ParsedSegment] = []
        # 按正文顺序遍历文档中的段落和表格块。
        for block in self._iter_docx_blocks(
            document,
            DocxDocument,
            CT_P,
            CT_Tbl,
            Paragraph,
            Table,
        ):
            # 当前块如果是普通段落，就直接按段落文本输出。
            if isinstance(block, Paragraph):
                # 跳过只有空白字符的段落。
                if block.text.strip():
                    segments.append(
                        ParsedSegment(
                            # 段落原文直接作为一个文本片段。
                            text=block.text,
                            # docx 当前不保留页码，因此设为 None。
                            page_start=None,
                            # docx 当前不保留结束页码，因此设为 None。
                            page_end=None,
                            # 记录来源类型是普通 docx 段落。
                            source_type="docx",
                        )
                    )
                # 段落处理完成后直接继续下一个块。
                continue

            # 当前块如果是表格，就逐行抽取表格内容。
            if isinstance(block, Table):
                # 遍历表格中的每一行。
                for row in block.rows:
                    # 把一整行单元格文本整理成一个语义片段。
                    row_text = self._build_docx_table_row_text(row.cells)
                    # 只有行文本非空时才输出片段。
                    if row_text:
                        segments.append(
                            ParsedSegment(
                                # 表格当前行的拼接文本作为片段正文。
                                text=row_text,
                                # docx 表格同样暂不保留页码。
                                page_start=None,
                                # docx 表格同样暂不保留结束页码。
                                page_end=None,
                                # 记录来源类型为表格行。
                                source_type="docx_table",
                            )
                        )
        # 返回统一的 ParsedDocument，供后续清洗与切块复用。
        return ParsedDocument(segments=segments)

    # 按 body 顺序遍历 docx 中的块级元素，避免只读 paragraphs 时漏掉表格正文。
    def _iter_docx_blocks(
        self,
        document,
        document_type,
        paragraph_type,
        table_type,
        paragraph_class,
        table_class,
    ):
        # 根据传入对象类型拿到真正的 body 节点，兼容文档根节点和子节点两种情况。
        parent = document.element.body if isinstance(document, document_type) else document._element
        # 按文档中的原始顺序遍历所有子节点。
        for child in parent.iterchildren():
            # 段落节点包装成高层 Paragraph 对象后 yield 给调用方。
            if isinstance(child, paragraph_type):
                yield paragraph_class(child, document)
            # 表格节点包装成高层 Table 对象后 yield 给调用方。
            elif isinstance(child, table_type):
                yield table_class(child, document)

    # 表格型文档把一行作为一个语义段，行内做轻量去重，避免合并单元格把同一文本重复很多次。
    def _build_docx_table_row_text(self, cells) -> str:
        # 保存当前行里按顺序出现的文本。
        texts: list[str] = []
        # 用集合记录已见过的单元格文本，避免重复追加。
        seen: set[str] = set()
        # 逐个处理当前行的单元格。
        for cell in cells:
            # 去掉单元格文本首尾空白。
            text = cell.text.strip()
            # 空文本或已出现过的文本都直接跳过。
            if not text or text in seen:
                continue
            # 把当前文本记入去重集合。
            seen.add(text)
            # 保留首次出现的单元格文本。
            texts.append(text)
        # 用竖线把一行里的多个单元格内容连接起来。
        return " | ".join(texts)

    # pdf 按页抽文本并保留页码，这样后面既能去页眉页脚，也能把检索结果定位回原页。
    def _parse_pdf(self, content: bytes) -> ParsedDocument:
        # 先尝试直接抽取每一页已有的文本层内容。
        page_texts = self._extract_pdf_page_texts(content)
        # 预留 OCR 结果字典，键是页索引，值是该页识别出的文本。
        ocr_texts: dict[int, str] = {}
        # 找出那些文本层为空的页索引，后面只对这些页做 OCR。
        missing_indexes = [
            index for index, text in enumerate(page_texts) if not text.strip()
        ]
        # 如果存在空白页，就启动 OCR 兜底。
        if missing_indexes:
            ocr_texts = self._ocr_pdf_page_texts(content, missing_indexes)

        # 保存最终输出的分页片段。
        segments = []
        # 枚举每一页文本，页码从 1 开始更符合用户认知。
        for index, extracted_text in enumerate(page_texts, start=1):
            # 默认优先使用原生提取到的文本。
            text = extracted_text
            # 默认来源类型标记为 pdf_text。
            source_type = "pdf_text"
            # 如果当前页原生文本为空，再尝试读取 OCR 结果。
            if not text.strip():
                # OCR 字典里存的是 0 基页索引，所以这里要减 1。
                text = ocr_texts.get(index - 1, "")
                # OCR 成功时标记为 pdf_ocr，否则仍记为 pdf_text。
                source_type = "pdf_ocr" if text.strip() else "pdf_text"
            # 把当前页输出成一个标准片段。
            segments.append(
                ParsedSegment(
                    # 当前页最终文本内容。
                    text=text,
                    # 起始页码就是当前页。
                    page_start=index,
                    # 结束页码同样是当前页。
                    page_end=index,
                    # 记录该页文本来源是原生文本还是 OCR。
                    source_type=source_type,
                )
            )
        # 返回包含所有页片段的标准文档对象。
        return ParsedDocument(segments=segments)

    # 原生文本提取优先走 pypdf，速度更快，且能保留文本层。
    def _extract_pdf_page_texts(self, content: bytes) -> list[str]:
        try:
            # PdfReader 负责读取 PDF 并访问每一页对象。
            from pypdf import PdfReader
        except ImportError as exc:
            # 没装 pypdf 时抛出明确异常，便于定位环境问题。
            raise RuntimeError("pypdf is not installed") from exc

        # 用内存字节流打开 PDF，避免先落临时文件。
        reader = PdfReader(BytesIO(content))
        # 逐页调用 extract_text，并把 None 兜底为空字符串。
        return [page.extract_text() or "" for page in reader.pages]

    # 对没有文本层的 PDF 页面走 OCR 兜底，适配扫描件和图片型 PDF。
    def _ocr_pdf_page_texts(self, content: bytes, page_indexes: list[int]) -> dict[int, str]:
        try:
            # fitz 负责把 PDF 页面渲染成位图图片。
            import fitz
        except ImportError:
            # 缺少渲染库时直接返回空结果，表示无法做 OCR 兜底。
            return {}

        try:
            # Pillow 用于把渲染得到的字节流打开成图像对象。
            from PIL import Image
        except ImportError:
            # 缺少图像库时也无法继续 OCR。
            return {}

        try:
            # pytesseract 负责真正执行 OCR 识别。
            import pytesseract
        except ImportError:
            # 缺少 OCR 引擎绑定时同样返回空结果。
            return {}

        # 在内存中打开 PDF 文档，准备逐页渲染。
        document = fitz.open(stream=content, filetype="pdf")
        # 用字典收集每个目标页识别出的文本。
        ocr_results: dict[int, str] = {}
        try:
            # 只对缺失文本层的页执行 OCR，减少不必要开销。
            for page_index in page_indexes:
                # 加载当前页对象。
                page = document.load_page(page_index)
                # 以 2 倍缩放渲染页面，提高 OCR 识别清晰度。
                pixmap = page.get_pixmap(alpha=False, matrix=fitz.Matrix(2, 2))
                # 把渲染结果转成 PNG 字节流，再交给 Pillow 打开。
                image = Image.open(BytesIO(pixmap.tobytes("png")))
                # 调用 Tesseract 做 OCR，并使用配置里的语言参数。
                ocr_results[page_index] = pytesseract.image_to_string(
                    image,
                    lang=self.settings.ocr_language,
                )
        finally:
            # 无论 OCR 过程是否出错，都关闭 PDF 文档句柄。
            document.close()

        # 返回只包含 OCR 过页面的文本字典。
        return ocr_results

    # png 先走 OCR，把图片统一转换成文本链路；当前阶段优先支持“图中文字检索”，不做图像向量。
    def _parse_png(self, content: bytes) -> ParsedDocument:
        try:
            # Pillow 用于把图片字节内容读成图像对象。
            from PIL import Image
        except ImportError as exc:
            # 缺少 Pillow 时抛出明确异常。
            raise RuntimeError("Pillow is not installed") from exc

        try:
            # pytesseract 负责从图片中识别文字。
            import pytesseract
        except ImportError as exc:
            # 缺少 OCR 库时同样抛出明确异常。
            raise RuntimeError("pytesseract is not installed") from exc

        # 把字节内容包装成图片对象。
        image = Image.open(BytesIO(content))
        # 按配置语言执行 OCR，得到整张图片的文本。
        text = pytesseract.image_to_string(image, lang=self.settings.ocr_language)
        # 把 OCR 文本封装成单页文档对象返回。
        return ParsedDocument(
            segments=[
                ParsedSegment(
                    # OCR 提取到的整图文本。
                    text=text,
                    # png 按单页处理，起始页码固定为 1。
                    page_start=1,
                    # png 按单页处理，结束页码固定为 1。
                    page_end=1,
                    # 记录来源类型为 OCR。
                    source_type="ocr",
                )
            ]
        )
