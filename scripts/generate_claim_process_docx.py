from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


DEFAULT_OUTPUT_PATH = Path("保险赔付流程规范（标准版）.docx")
APPENDIX_OUTPUT_PATH = Path("保险赔付流程规范（章节说明版）.docx")


PROCESS_HEADERS = [
    "节点ID",
    "节点名称",
    "节点类型",
    "执行角色",
    "前置节点",
    "输出节点",
    "输入资料",
    "输出资料",
    "处理规则",
    "异常去向",
    "表单编码",
]

PROCESS_ROWS = [
    [
        "N001",
        "开始报案",
        "start",
        "客户",
        "-",
        "N002",
        "保单号;出险时间;事故描述;报案人信息",
        "报案申请",
        "客户发起赔付申请时创建理赔单号",
        "-",
        "claim_report_form",
    ],
    [
        "N002",
        "资料提交",
        "userTask",
        "客户",
        "N001",
        "N003",
        "身份证明;保单信息;医疗/事故证明;收款账户",
        "索赔资料包",
        "客户需一次性上传基础索赔资料",
        "资料缺失时由N004触发补件",
        "claim_report_form",
    ],
    [
        "N003",
        "资料初审",
        "userTask",
        "理赔专员",
        "N002",
        "N004",
        "索赔资料包",
        "初审结论",
        "检查资料完整性、清晰度、真实性基础项",
        "可转N005补件通知",
        "document_review_form",
    ],
    [
        "N004",
        "是否资料齐全",
        "exclusiveGateway",
        "系统",
        "N003",
        "N005|N006",
        "初审结论",
        "分支结果",
        "若资料缺失则流转补件，否则进入受理",
        "缺失->N005;齐全->N006",
        "-",
    ],
    [
        "N005",
        "补件通知",
        "userTask",
        "理赔专员",
        "N004",
        "N002",
        "缺失项清单",
        "补件通知单",
        "明确列出需补交资料及截止时间",
        "客户补件后回到N002",
        "supplement_notice_form",
    ],
    [
        "N006",
        "理赔受理",
        "userTask",
        "理赔专员",
        "N004",
        "N007",
        "完整索赔资料包",
        "受理结果",
        "生成受理记录并锁定受理时间",
        "疑似异常案件可备注进入人工复核",
        "acceptance_form",
    ],
    [
        "N007",
        "责任审核",
        "userTask",
        "审核岗",
        "N006",
        "N008",
        "保单条款;事故材料;受理结果",
        "责任审核意见",
        "核验是否属于保障责任及免责范围",
        "不属于责任时进入拒赔",
        "coverage_review_form",
    ],
    [
        "N008",
        "是否属于保险责任",
        "exclusiveGateway",
        "系统",
        "N007",
        "N009|N013",
        "责任审核意见",
        "分支结果",
        "属于保险责任则进入金额核算，否则拒赔结案",
        "是->N009;否->N013",
        "-",
    ],
    [
        "N009",
        "金额核算",
        "userTask",
        "理算岗",
        "N008",
        "N010",
        "责任审核意见;费用明细;保额信息",
        "赔付测算单",
        "按条款、免赔额、赔付比例计算应赔金额",
        "大额案件可进入人工复核",
        "amount_calc_form",
    ],
    [
        "N010",
        "人工复核",
        "userTask",
        "复核岗",
        "N009",
        "N011",
        "赔付测算单;高风险标记",
        "复核意见",
        "金额超阈值或命中风险规则时必须复核",
        "复核不通过可退回N009重算",
        "manual_review_form",
    ],
    [
        "N011",
        "赔付审批",
        "userTask",
        "审批主管",
        "N010",
        "N012",
        "复核意见;赔付测算单",
        "审批结果",
        "审批通过后进入付款执行",
        "审批驳回可退回N009",
        "approval_form",
    ],
    [
        "N012",
        "付款执行",
        "serviceTask",
        "财务系统",
        "N011",
        "N014",
        "审批结果;收款账户",
        "付款回执",
        "调用付款接口并记录支付流水",
        "支付失败转人工处理",
        "payment_form",
    ],
    [
        "N013",
        "拒赔结案通知",
        "endTask",
        "理赔专员",
        "N008",
        "N015",
        "责任审核意见",
        "拒赔通知书",
        "说明拒赔原因及申诉渠道",
        "-",
        "rejection_notice_form",
    ],
    [
        "N014",
        "结案通知",
        "userTask",
        "理赔专员",
        "N012",
        "N015",
        "付款回执",
        "结案通知书",
        "通知客户赔付完成并同步支付信息",
        "-",
        "closure_notice_form",
    ],
    [
        "N015",
        "归档结束",
        "end",
        "系统",
        "N013|N014",
        "-",
        "通知结果;理赔资料",
        "归档记录",
        "归档所有流程记录并结束案件",
        "-",
        "archive_form",
    ],
]

FORM_HEADERS = [
    "表单编码",
    "字段编码",
    "字段名称",
    "字段类型",
    "是否必填",
    "默认值",
    "枚举值",
    "校验规则",
    "回填来源",
]

FORM_ROWS = [
    ["claim_report_form", "policy_no", "保单号", "string", "Y", "", "", "长度>=8", "用户输入"],
    ["claim_report_form", "claim_no", "理赔单号", "string", "Y", "系统生成", "", "唯一值", "系统生成"],
    ["claim_report_form", "incident_time", "出险时间", "datetime", "Y", "", "", "不得晚于当前时间", "用户输入"],
    ["claim_report_form", "claimant_name", "报案人姓名", "string", "Y", "", "", "非空", "用户输入"],
    ["document_review_form", "doc_complete", "资料是否齐全", "boolean", "Y", "", "true,false", "必填", "N003审核结果"],
    ["document_review_form", "missing_items", "缺失项清单", "textarea", "N", "", "", "资料不齐全时必填", "N003审核结果"],
    ["supplement_notice_form", "supplement_deadline", "补件截止日", "date", "Y", "", "", "晚于通知日", "人工填写"],
    ["supplement_notice_form", "supplement_reason", "补件原因", "textarea", "Y", "", "", "非空", "N005填写"],
    ["acceptance_form", "acceptance_time", "受理时间", "datetime", "Y", "当前时间", "", "自动写入", "系统生成"],
    ["acceptance_form", "case_level", "案件等级", "select", "Y", "standard", "standard,priority,risk", "枚举值", "N006判断"],
    ["coverage_review_form", "coverage_result", "责任判定", "select", "Y", "", "covered,excluded,uncertain", "枚举值", "N007审核结果"],
    ["coverage_review_form", "coverage_comment", "责任说明", "textarea", "Y", "", "", "非空", "N007审核结果"],
    ["amount_calc_form", "claim_amount", "申请金额", "number", "Y", "", "", ">=0", "资料包提取"],
    ["amount_calc_form", "approved_amount", "核赔金额", "number", "Y", "", "", ">=0", "N009计算结果"],
    ["amount_calc_form", "deductible_amount", "免赔额", "number", "N", "0", "", ">=0", "条款计算"],
    ["manual_review_form", "risk_flag", "风险标记", "multiSelect", "Y", "", "high_amount,duplicate_claim,suspected_fraud", "至少1项", "规则命中"],
    ["manual_review_form", "review_result", "复核结论", "select", "Y", "", "pass,rework,reject", "枚举值", "N010复核结果"],
    ["approval_form", "approval_result", "审批结果", "select", "Y", "", "approved,rejected", "枚举值", "N011审批结果"],
    ["approval_form", "approval_comment", "审批意见", "textarea", "N", "", "", "驳回时必填", "N011审批结果"],
    ["payment_form", "payee_account", "收款账户", "string", "Y", "", "", "银行卡号或账户号格式校验", "资料包提取"],
    ["payment_form", "payment_status", "付款状态", "select", "Y", "pending", "pending,success,failed", "枚举值", "N012执行结果"],
    ["rejection_notice_form", "rejection_reason", "拒赔原因", "textarea", "Y", "", "", "非空", "N007责任审核意见"],
    ["closure_notice_form", "payment_reference", "支付流水号", "string", "Y", "", "", "非空", "N012付款回执"],
    ["archive_form", "archive_status", "归档状态", "select", "Y", "archived", "archived", "固定值", "系统生成"],
]

EDGE_HEADERS = [
    "边ID",
    "起始节点ID",
    "目标节点ID",
    "条件类型",
    "条件表达式",
    "说明",
]

EDGE_ROWS = [
    ["E001", "N001", "N002", "always", "", "客户报案后进入资料提交"],
    ["E002", "N002", "N003", "always", "", "资料提交后进入初审"],
    ["E003", "N003", "N004", "always", "", "初审结束进入资料完整性判断"],
    ["E004", "N004", "N005", "expression", "doc_complete == false", "资料不齐全时通知补件"],
    ["E005", "N005", "N002", "always", "", "客户补件后重新提交资料"],
    ["E006", "N004", "N006", "expression", "doc_complete == true", "资料齐全时进入受理"],
    ["E007", "N006", "N007", "always", "", "受理后进入责任审核"],
    ["E008", "N007", "N008", "always", "", "责任审核结束进入责任判断"],
    ["E009", "N008", "N009", "expression", "coverage_result == 'covered'", "属于保险责任时进入金额核算"],
    ["E010", "N008", "N013", "expression", "coverage_result != 'covered'", "不属于保险责任时拒赔结案"],
    ["E011", "N009", "N010", "expression", "approved_amount >= 10000 or risk_flag != ''", "大额或高风险案件进入人工复核"],
    ["E012", "N010", "N009", "expression", "review_result == 'rework'", "复核要求重算时退回金额核算"],
    ["E013", "N010", "N011", "expression", "review_result == 'pass'", "复核通过进入审批"],
    ["E014", "N009", "N011", "expression", "approved_amount < 10000 and risk_flag == ''", "普通案件可直接审批"],
    ["E015", "N011", "N012", "expression", "approval_result == 'approved'", "审批通过进入付款"],
    ["E016", "N011", "N009", "expression", "approval_result == 'rejected'", "审批驳回退回金额核算"],
    ["E017", "N012", "N014", "expression", "payment_status == 'success'", "付款成功后结案通知"],
    ["E018", "N013", "N015", "always", "", "拒赔通知后归档"],
    ["E019", "N014", "N015", "always", "", "结案通知后归档"],
]

EXCEPTION_HEADERS = ["规则编码", "场景", "触发条件", "处理动作", "回退节点", "备注"]

EXCEPTION_ROWS = [
    ["R001", "补件", "资料不齐全", "发送补件通知并挂起案件", "N002", "补件后重新进入资料提交流程"],
    ["R002", "拒赔", "不属于保险责任", "生成拒赔通知书并结束案件", "N013", "记录拒赔原因与申诉渠道"],
    ["R003", "人工复核", "高金额或命中风险规则", "转复核岗处理", "N010", "复核通过后继续审批"],
    ["R004", "审批驳回", "审批不通过", "退回理算岗重新核算", "N009", "需补充审批意见"],
    ["R005", "支付失败", "付款接口返回失败", "转人工处理并补录结果", "N012", "标准版样例中不展开支付失败支线"],
]


def set_cell_text(cell, text: str) -> None:
    cell.text = str(text)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(9)


def shade_header_cell(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "D9EAF7")
    tc_pr.append(shd)


def add_table(document: Document, title: str, headers: list[str], rows: list[list[str]]) -> None:
    document.add_heading(title, level=2)
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        set_cell_text(header_cells[index], header)
        shade_header_cell(header_cells[index])

    for row in rows:
        row_cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(row_cells[index], value)

    document.add_paragraph("")


def add_cover(
    document: Document,
    title_text: str = "保险赔付流程规范（标准版）",
    subtitle_text: str = "机器可解析流程文件样例",
    note_text: str = "本样例优先服务于自动化解析，采用“流程主表 + 节点表单定义 + 连线规则 + 异常规则”的固定结构。",
) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(title_text)
    run.bold = True
    run.font.size = Pt(20)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subrun = subtitle.add_run(subtitle_text)
    subrun.font.size = Pt(12)

    document.add_paragraph("")
    meta = document.add_table(rows=4, cols=2)
    meta.style = "Table Grid"
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    metadata = [
        ("流程编号", "CLAIM-STD-001"),
        ("流程名称", "保险赔付标准流程"),
        ("版本", "v1.0"),
        ("适用范围", "通用个人保险赔付受理、审核、核赔、审批、付款"),
    ]
    for row_index, (key, value) in enumerate(metadata):
        set_cell_text(meta.rows[row_index].cells[0], key)
        shade_header_cell(meta.rows[row_index].cells[0])
        set_cell_text(meta.rows[row_index].cells[1], value)

    document.add_paragraph("")
    note = document.add_paragraph()
    note.style = document.styles["Normal"]
    note.add_run("生成说明：").bold = True
    note.add_run(note_text)


def add_usage_notes(document: Document) -> None:
    document.add_heading("一、解析约束说明", level=1)
    bullets = [
        "节点ID在全文唯一，建议作为画布 XML 节点主键。",
        "节点类型仅允许使用 start、end、userTask、serviceTask、exclusiveGateway、endTask。",
        "前置节点和输出节点可包含多个值，多个值之间用英文竖线分隔。",
        "表单编码用于节点与表单模板绑定，字段值回填来源在“节点表单定义”中声明。",
        "条件表达式使用简化布尔表达式，适合作为 LangGraph 中间结构的分支条件。",
    ]
    for text in bullets:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(text)


def add_narrative_sections(document: Document) -> None:
    document.add_heading("一、文档说明", level=1)
    paragraphs = [
        "本规范用于描述个人保险赔付的标准处理路径，覆盖报案、资料提交、资料初审、责任审核、金额核算、审批付款与结案归档等核心环节。",
        "本版本采用“章节说明 + 结构化附录”的组织方式。正文面向业务人员阅读，附录面向系统解析与画布生成。",
        "当业务人员阅读正文时，可快速理解各环节职责与处理要点；当系统读取附录时，可直接抽取节点、表单和连线信息生成流程图与节点表单。",
    ]
    for text in paragraphs:
        document.add_paragraph(text)

    document.add_heading("二、流程概述", level=1)
    overview = [
        ("2.1 适用范围", "适用于标准个人保险赔付案件，不包含复杂争议案件、重大团险赔付及跨机构联动赔案。"),
        ("2.2 参与角色", "主要参与角色包括客户、理赔专员、审核岗、理算岗、复核岗、审批主管及财务系统。"),
        ("2.3 主流程", "标准路径为：开始报案 -> 资料提交 -> 资料初审 -> 资料完整性判断 -> 理赔受理 -> 责任审核 -> 责任判定 -> 金额核算 -> 人工复核/审批 -> 付款执行 -> 结案通知 -> 归档结束。"),
    ]
    for title, text in overview:
        document.add_paragraph().add_run(title).bold = True
        document.add_paragraph(text)

    document.add_heading("三、节点说明", level=1)
    node_descriptions = [
        ("3.1 开始报案与资料提交", "客户发起赔付申请后，系统创建理赔单号。客户需按要求上传保单信息、事故材料、身份证明及收款账户等基础资料。"),
        ("3.2 资料初审与补件处理", "理赔专员对资料完整性、清晰度和基础真实性进行检查。若存在缺失项，应生成补件通知并要求客户重新提交资料。"),
        ("3.3 理赔受理与责任审核", "资料齐全后由理赔专员正式受理，审核岗根据保单条款和事故材料判断是否属于保险责任范围。"),
        ("3.4 金额核算与人工复核", "责任成立的案件进入理算环节，根据保额、免赔额和赔付比例测算应赔金额。高金额或高风险案件需转复核岗人工复核。"),
        ("3.5 审批、付款与结案", "审批主管根据测算单和复核意见作出审批决定。审批通过后由财务系统付款，付款成功后通知客户并归档结案。"),
    ]
    for title, text in node_descriptions:
        document.add_paragraph().add_run(title).bold = True
        document.add_paragraph(text)

    document.add_heading("四、异常处理原则", level=1)
    exception_paragraphs = [
        "资料不齐全时，必须发送补件通知并回退到资料提交节点，不允许直接跳过补件进入受理。",
        "责任审核不通过时，应生成拒赔通知书并告知客户拒赔原因与申诉渠道。",
        "金额超过阈值或命中风险规则时，必须进入人工复核节点，由复核岗给出结论。",
        "审批驳回时，案件退回理算岗重新核算，并补充审批意见。",
    ]
    for text in exception_paragraphs:
        document.add_paragraph(style="List Bullet").add_run(text)


def apply_base_style(document: Document) -> Document:
    styles = document.styles
    styles["Normal"].font.name = "Songti SC"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(10.5)
    return document


def build_machine_first_document() -> Document:
    document = apply_base_style(Document())
    add_cover(document)
    document.add_section(WD_SECTION.NEW_PAGE)
    add_usage_notes(document)
    add_table(document, "二、流程主表", PROCESS_HEADERS, PROCESS_ROWS)
    add_table(document, "三、节点表单定义", FORM_HEADERS, FORM_ROWS)
    add_table(document, "四、连线规则表", EDGE_HEADERS, EDGE_ROWS)
    add_table(document, "五、状态与异常规则", EXCEPTION_HEADERS, EXCEPTION_ROWS)
    return document


def build_appendix_document() -> Document:
    document = apply_base_style(Document())
    add_cover(
        document,
        title_text="保险赔付流程规范（章节说明版）",
        subtitle_text="章节说明 + 结构化附录样例",
        note_text="本样例在正文保留业务阅读性，并在附录中提供结构化节点、表单和连线定义，兼顾人工阅读与系统解析。",
    )
    document.add_section(WD_SECTION.NEW_PAGE)
    add_narrative_sections(document)
    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_heading("附录A 结构化流程主表", level=1)
    add_table(document, "A.1 节点清单", PROCESS_HEADERS, PROCESS_ROWS)
    document.add_heading("附录B 节点表单定义", level=1)
    add_table(document, "B.1 表单字段清单", FORM_HEADERS, FORM_ROWS)
    document.add_heading("附录C 连线与异常规则", level=1)
    add_table(document, "C.1 连线规则表", EDGE_HEADERS, EDGE_ROWS)
    add_table(document, "C.2 状态与异常规则", EXCEPTION_HEADERS, EXCEPTION_ROWS)
    return document


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--mode",
        choices=["machine", "appendix"],
        default="machine",
        help="machine: 机器可解析优先；appendix: 章节说明 + 结构化附录",
    )
    args = parser.parse_args()

    if args.mode == "appendix":
        document = build_appendix_document()
        output_path = APPENDIX_OUTPUT_PATH
    else:
        document = build_machine_first_document()
        output_path = DEFAULT_OUTPUT_PATH

    document.save(output_path)
    print(output_path.resolve())


if __name__ == "__main__":
    main()
